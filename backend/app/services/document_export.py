import os
import json
import logging
from datetime import datetime
from typing import Optional
from app.core.config import settings

logger = logging.getLogger(__name__)


class DocumentExportService:
    def __init__(self):
        self.export_dir = os.path.join(settings.storage_local_path, "exports")
        os.makedirs(self.export_dir, exist_ok=True)

    def process_export(self, export_id: str) -> dict:
        from app.core.database import async_session_factory
        from app.models.import_export import ExportRecord
        from sqlalchemy import select

        async def _process():
            async with async_session_factory() as db:
                result = await db.execute(
                    select(ExportRecord).where(ExportRecord.id == export_id)
                )
                record = result.scalar_one_or_none()
                if not record:
                    return {"error": "Export not found"}

                try:
                    if record.format == "pdf":
                        file_path = self._export_pdf(record)
                    elif record.format == "docx":
                        file_path = self._export_docx(record)
                    elif record.format == "markdown":
                        file_path = self._export_markdown(record)
                    elif record.format == "html":
                        file_path = self._export_html(record)
                    else:
                        record.status = "failed"
                        record.error_message = f"Unsupported format: {record.format}"
                        return {"error": record.error_message}

                    record.status = "completed"
                    record.file_path = file_path
                    record.completed_at = datetime.utcnow()
                    if os.path.exists(file_path):
                        record.file_size = os.path.getsize(file_path)
                    return {"status": "completed", "file_path": file_path}

                except Exception as e:
                    record.status = "failed"
                    record.error_message = str(e)
                    logger.exception(f"Export failed: {e}")
                    return {"error": str(e)}

        import asyncio
        return asyncio.run(_process())

    def _export_pdf(self, record) -> str:
        from weasyprint import HTML
        html = self._generate_html(record)
        filename = f"{record.id}.pdf"
        filepath = os.path.join(self.export_dir, filename)
        HTML(string=html).write_pdf(filepath)
        return filepath

    def _export_docx(self, record) -> str:
        from docx import Document
        doc = Document()
        resume_data = record.options or {}
        doc.add_heading(resume_data.get("title", "Resume"), 0)
        doc.add_paragraph(resume_data.get("summary", ""))
        for exp in resume_data.get("experiences", []):
            doc.add_heading(f"{exp.get('position', '')} at {exp.get('company', '')}", level=2)
            doc.add_paragraph(exp.get("description", ""))
        filename = f"{record.id}.docx"
        filepath = os.path.join(self.export_dir, filename)
        doc.save(filepath)
        return filepath

    def _export_markdown(self, record) -> str:
        resume_data = record.options or {}
        lines = []
        lines.append(f"# {resume_data.get('title', 'Resume')}")
        lines.append("")
        summary = resume_data.get("summary", "")
        if summary:
            lines.append(summary)
            lines.append("")
        for exp in resume_data.get("experiences", []):
            lines.append(f"## {exp.get('position', '')} at {exp.get('company', '')}")
            lines.append("")
            desc = exp.get("description", "")
            if desc:
                lines.append(desc)
                lines.append("")
        filename = f"{record.id}.md"
        filepath = os.path.join(self.export_dir, filename)
        with open(filepath, "w") as f:
            f.write("\n".join(lines))
        return filepath

    def _export_html(self, record) -> str:
        resume_data = record.options or {}
        title = resume_data.get("title", "Resume")
        summary = resume_data.get("summary", "")
        experiences_html = ""
        for exp in resume_data.get("experiences", []):
            experiences_html += f"""
            <div class="experience">
                <h2>{exp.get('position', '')} at {exp.get('company', '')}</h2>
                <p>{exp.get('description', '')}</p>
            </div>
            """
        html = f"""<!DOCTYPE html>
<html lang="en">
<head><meta charset="UTF-8"><title>{title}</title>
<style>
body {{ font-family: 'Inter', Arial, sans-serif; max-width: 800px; margin: 40px auto; padding: 20px; }}
h1 {{ font-size: 28px; margin-bottom: 4px; }}
.experience {{ margin-bottom: 20px; }}
</style>
</head>
<body>
<h1>{title}</h1>
<p>{summary}</p>
{experiences_html}
</body>
</html>"""
        filename = f"{record.id}.html"
        filepath = os.path.join(self.export_dir, filename)
        with open(filepath, "w") as f:
            f.write(html)
        return filepath

    def _generate_html(self, record) -> str:
        resume_data = record.options or {}
        title = resume_data.get("title", "Resume")
        pi = resume_data.get("personal_information", {})
        summary = resume_data.get("summary", "")

        experiences_html = ""
        for exp in resume_data.get("experiences", []):
            techs = " ".join(
                f"<span class='tech'>{t}</span>" for t in exp.get("technologies", [])
            )
            achievements = "".join(
                f"<li>{a}</li>" for a in exp.get("achievements", [])
            )
            experiences_html += f"""
            <div class="experience">
                <div class="exp-header">
                    <strong>{exp.get('position', '')}</strong> at {exp.get('company', '')}
                </div>
                <p>{exp.get('description', '')}</p>
                {techs}
                <ul>{achievements}</ul>
            </div>
            """

        education_html = ""
        for edu in resume_data.get("education", []):
            education_html += f"<p><strong>{edu.get('degree', '')}</strong> - {edu.get('institution', '')}</p>"

        skills_html = ""
        for skill in resume_data.get("skills", []):
            skills_html += f"<span class='skill'>{skill.get('skill_name', '')}</span> "

        html = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<title>{pi.get('full_name', title)} - Resume</title>
<style>
@page {{ size: A4; margin: 20mm; }}
body {{ font-family: 'Inter', 'Segoe UI', Arial, sans-serif; font-size: 11pt; line-height: 1.5; color: #333; }}
h1 {{ font-size: 24pt; margin-bottom: 2px; }}
h2 {{ font-size: 14pt; border-bottom: 2px solid #2563eb; padding-bottom: 4px; margin-top: 20px; }}
.contact {{ color: #666; font-size: 10pt; margin-bottom: 16px; }}
.summary {{ margin: 12px 0; }}
.experience {{ margin-bottom: 16px; }}
.exp-header {{ font-weight: 600; margin-bottom: 4px; }}
.tech {{ display: inline-block; background: #e8f0fe; color: #2563eb; padding: 1px 8px; border-radius: 3px; font-size: 9pt; margin: 2px; }}
.skill {{ display: inline-block; background: #f3f4f6; padding: 2px 10px; border-radius: 3px; margin: 2px; font-size: 10pt; }}
ul {{ margin: 4px 0; padding-left: 20px; }}
</style>
</head>
<body>
<h1>{pi.get('full_name', '')}</h1>
<div class="contact">
{pi.get('email', '')} | {pi.get('phone', '')} | {pi.get('city', '')} {pi.get('state', '')}
</div>
<div class="summary"><p>{summary}</p></div>
<h2>Experience</h2>
{experiences_html}
<h2>Education</h2>
{education_html}
<h2>Skills</h2>
<p>{skills_html}</p>
</body>
</html>"""
        return html
